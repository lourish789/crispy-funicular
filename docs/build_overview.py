"""Build the SHORT AgriTech Suite overview (.docx) — a quick picture of the idea.

Target: 5-6 pages. Narrative, not a business proposal. No contents page, no
executive summary. Strictly monochrome.
"""
from __future__ import annotations

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
BLACK = RGBColor(0, 0, 0)
OUT = os.path.join(HERE, "AgriTech_Suite_Overview.docx")

doc = Document()
for s in doc.sections:
    s.top_margin = Inches(0.8)
    s.bottom_margin = Inches(0.8)
    s.left_margin = Inches(0.9)
    s.right_margin = Inches(0.9)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = BLACK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1

for nm, sz in [("Title", 22), ("Heading 1", 13.5), ("Heading 2", 11.5)]:
    st = doc.styles[nm]
    st.font.color.rgb = BLACK
    st.font.name = "Calibri"
    st.font.size = Pt(sz)
    st.paragraph_format.space_before = Pt(10)
    st.paragraph_format.space_after = Pt(4)


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = BLACK
    return p


def para(text, italic=False, size=10.5, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = BLACK
    if align is not None:
        p.alignment = align
    return p


def _shade(cell, hexval="D9D9D9"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexval)
    tcPr.append(shd)


def table(headers, rows, widths=None, font=9):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(htext)
        r.bold = True
        r.font.size = Pt(font)
        r.font.color.rgb = BLACK
        _shade(c)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val))
            r.font.size = Pt(font)
            r.font.color.rgb = BLACK
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def figure(name, caption, width=5.5):
    doc.add_picture(os.path.join(HERE, name), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption)
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = BLACK


# ---------------------------------------------------------------- header
p = doc.add_paragraph()
r = p.add_run("AGRITECH SUITE")
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = BLACK
p.paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
r = p.add_run("One platform for the African farmer: diagnose, decide, sell.")
r.font.size = Pt(12)
r.italic = True
r.font.color.rgb = BLACK
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
r = p.add_run("_" * 96)
r.font.color.rgb = BLACK
p.paragraph_format.space_after = Pt(8)

# ---------------------------------------------------------------- idea
h("The idea", 1)
para(
    "A farmer in Nigeria wakes up to find leaves turning brown. Nobody is coming to tell her "
    "what it is. The nearest extension officer serves thousands of other farmers and she has "
    "never met him. She guesses, sprays whatever the shop recommends, and hopes. Weeks later, "
    "whatever survives is sold at the farm gate to a middleman, at whatever price he names, "
    "because she has no idea what the crop is worth in Lagos and no way to reach anyone there."
)
para(
    "AgriTech Suite is built for that morning. She photographs the leaf, and in seconds the "
    "app names the disease, explains what caused it, tells her what to do today and how to stop "
    "it next season. The same app tells her what to grow more of and what to abandon, based on "
    "her own farm's history. When she harvests, it shows her buyers who are looking for exactly "
    "what she has — and shows those buyers her produce, directly."
)
para(
    "It is one platform doing six jobs that are currently scattered across six apps, most of "
    "which the average smallholder will never install: AI disease diagnosis for crops and "
    "livestock, personalised farm advisory, local agricultural news, a farmer-to-buyer "
    "marketplace, a community of peers, and a conversational AI expert that already knows her "
    "farm."
)

# ---------------------------------------------------------------- problem
h("The problem this exists to solve", 1)
para(
    "African agriculture is not short of effort. It is short of information reaching the person "
    "holding the hoe, and short of a fair route to market once the harvest is in. Three numbers "
    "explain the whole business case."
)
table(
    ["The number", "What it means"],
    [
        ["1 extension officer per 2,500-10,000 farmers\n(FAO recommends 1 per 1,000; some Nigerian states report 1 per 24,000)",
         "Expert advice effectively never arrives. Most farmers will not meet an extension officer this season. When disease strikes, they guess."],
        ["N3.5-5 trillion lost in Nigeria in 2025\n(30-40 million tonnes of food; up to 50% for tomatoes, onions, bananas)",
         "Farmers do everything right in the field and still lose the value after harvest — to disease caught too late, and to having no buyer ready."],
        ["145 agritech companies in Nigeria — yet thin smallholder adoption",
         "The tools exist but do not fit: single-purpose, priced for commercial farms, or requiring data and smartphones that rural users do not have."],
    ],
    widths=[2.5, 4.1],
)
para(
    "These compound. Late diagnosis destroys the crop; no market access destroys the price of "
    "whatever survived. Fix one and you have helped a little. Fix both in the same place and "
    "the farmer's season changes."
)

# ---------------------------------------------------------------- who
h("Who it is for", 1)
para(
    "Nigerian farmers and agricultural traders first, then West Africa and the wider continent. "
    "Nothing in the system is Nigeria-specific — crops, languages, currencies and news sources "
    "are all configurable. Every account is one of two roles, and that single choice shapes the "
    "entire experience."
)
table(
    ["Role", "Who they are", "What they get"],
    [
        ["Farmer (Producer)",
         "Smallholders, mixed crop-and-livestock households, commercial growers, cooperatives.",
         "Disease diagnosis, personalised advisory, ability to list produce, and a feed of buyers currently seeking what they grow."],
        ["Buyer (Consumer)",
         "Market traders, aggregators, processors, restaurants, exporters, bulk households.",
         "Ability to post buy requests, and a feed of produce offered directly by farmers, with location and price."],
    ],
    widths=[1.1, 2.5, 3.0],
)

# ---------------------------------------------------------------- how
h("How it works", 1)
para(
    "The most-used moment is the diagnosis. It was deliberately built so that it still returns "
    "a useful answer when the network is weak or absent — because that is exactly where farmers "
    "are standing when they need it."
)
figure("fig2_diagnosis_flow.png",
       "The diagnosis flow. If anything in the chain is unavailable, it falls back rather than failing.")
para(
    "The commercial engine is the marketplace, and it is two-sided by design. A farmer never has "
    "to understand listing categories — they post what they have, and the system publishes it as "
    "a sell offer. A buyer posts what they need and it becomes a buy request. Each side is then "
    "shown the other."
)
figure("fig3_marketplace.png", "Producers and buyers are matched directly, cutting out the intermediary.")

# ---------------------------------------------------------------- different
h("What makes it different", 1)
para(
    "The market has good point solutions. PlantVillage Nuru diagnoses crop disease offline and "
    "does it well. Crop2Cash reaches feature-phone farmers with financial services. AFEX runs "
    "real commodity trading. Farmcrowdy and Thrive Agric bring financing. None of them assemble "
    "the whole picture for one farmer in one place."
)
table(
    ["", "Nuru", "Plantix", "Farmcrowdy", "Thrive Agric", "AFEX", "Crop2Cash", "AgriTech Suite"],
    [
        ["Plant diagnosis", "Yes", "Yes", "No", "Partial", "No", "No", "Yes"],
        ["Animal diagnosis", "No", "No", "No", "No", "No", "No", "Yes"],
        ["Per-farm advisory", "Partial", "No", "No", "Partial", "No", "No", "Yes"],
        ["Direct farmer-to-buyer market", "No", "No", "Partial", "Partial", "Yes", "No", "Yes"],
        ["AI assistant", "No", "No", "No", "No", "No", "No", "Yes"],
        ["Free to run (no paid AI per use)", "Yes", "No", "n/a", "n/a", "n/a", "n/a", "Yes"],
        ["Live in the field today", "Yes", "Yes", "Yes", "Yes", "Yes", "Yes", "Not yet"],
    ],
    widths=[1.75, 0.62, 0.68, 0.85, 0.85, 0.55, 0.78, 0.92],
    font=8.5,
)
para(
    "Two honest conclusions. First, nobody else combines diagnosis, advisory and a direct "
    "marketplace so that the diagnosis informs what to plant and where to sell it — that gap is "
    "the business. Second, this project is the only one in that table not yet in the field. It "
    "is technically complete and unproven, and the plan reflects that."
)
para(
    "The quiet advantage is technical but has a commercial consequence: the disease model is "
    "open-source and runs locally rather than calling a paid AI service. That means each "
    "diagnosis costs effectively nothing to serve. It is the difference between a tool a "
    "subsistence farmer can use freely forever and one that must be metered — and metering is "
    "precisely what prices out the farmers who need it most."
)

# ---------------------------------------------------------------- inclusivity
h("It has to work for the poorest farmer, not just the richest", 1)
para(
    "Users will not be alike. Some will run hundreds of hectares from a laptop. Far more will "
    "have a shared phone, a weak signal, a prepaid balance measured in hundreds of naira, and "
    "limited reading confidence in English. A platform that only serves the first group is not "
    "an agricultural platform — it is a business tool for people who already have options."
)
para(
    "So capability degrades with the user's means but never disappears. The core — diagnosis, "
    "advisory, news, community, AI assistant — is free permanently, and revenue never comes from "
    "access to knowledge. A planned USSD short-code and SMS layer will carry price alerts, "
    "buyer matching and advisory summaries to feature phones with no data bundle at all, and "
    "local languages (Hausa, Yoruba, Igbo, Pidgin) with voice answers will follow for users who "
    "cannot comfortably read English. Where the device cannot reach, a trusted local agent will "
    "— earning a commission, which turns inclusion itself into a livelihood."
)

# ---------------------------------------------------------------- money
h("How it makes money", 1)
para(
    "One rule: the platform earns when the user earns, and not before. Charge a subsistence "
    "farmer for advice and they will abandon the app, and they would be right to. Take a small "
    "share of a sale they would not otherwise have made, and they will gladly pay it."
)
figure("fig6_revenue.png", "Revenue sits on top of a permanently free core.", width=5.3)
table(
    ["Stream", "Who pays", "Why it works"],
    [
        ["Marketplace commission (primary)", "Buyer or seller, on a completed trade",
         "Perfectly aligned: no sale, no fee. Nothing is paid upfront and no behaviour has to change. This switches on first."],
        ["Premium subscription", "Active traders, commercial farmers",
         "Verified badge, priority listing, price alerts, better matching. Optional — never required for the core."],
        ["Business / API tier", "Cooperatives, processors, exporters, NGOs",
         "Bulk sourcing, analytics, integrations. These users have budgets and clear operational value."],
        ["Input & service referrals", "Seed, fertiliser, veterinary, logistics providers",
         "The diagnosis already identified the need; connecting it to a vetted supplier serves the farmer and monetises the moment."],
        ["Financing & insurance partnerships", "Lenders, insurers",
         "Verified sales history de-risks lending — and gives farmers a credit record they have never had."],
    ],
    widths=[1.6, 1.8, 3.2],
)
para(
    "The farmer's side of the ledger matters just as much: a better price by reaching buyers "
    "directly, less of the harvest lost to late diagnosis, money not wasted on the wrong "
    "chemical, planting what is actually in demand, and eventually a financial record that "
    "unlocks credit. The platform's revenue is a by-product of those gains, which is the only "
    "arrangement that survives contact with a farmer's budget."
)

# ---------------------------------------------------------------- status
h("Where it stands", 1)
para(
    "The software is built and tested — not a prototype or a mock-up. The backend, the web "
    "application, the AI pipeline, the role system and the database all work, and an automated "
    "test suite covering every feature passes."
)
table(
    ["Built and working", "Not yet done"],
    [
        ["Disease diagnosis for plants (38 classes) and livestock, from photo or video",
         "Deployment to public hosting (Render for the backend, Vercel for the frontend)"],
        ["Personalised advisory, agri-news, marketplace, community, streaming AI assistant",
         "A field pilot with real farmers — accuracy in real conditions is not yet measured"],
        ["Farmer/Buyer roles, sign-in (including Firebase), full database and web interface",
         "USSD and SMS access, local languages, payments and escrow"],
    ],
    widths=[3.3, 3.3],
)
para(
    "The immediate next step is simply to put it online and get it in front of real farmers in "
    "one Nigerian state. The remaining work is about reach and evidence rather than "
    "construction: deploy it, measure the diagnosis honestly against expert assessment, then "
    "build the USSD layer that extends it to the feature-phone majority that no smartphone-only "
    "platform will ever serve."
)

# ---------------------------------------------------------------- contact
h("Contact", 1)
t = doc.add_table(rows=0, cols=2)
t.style = "Table Grid"
for k, v in [
    ("Author", "Flourish Olaiya — Full-Stack Software Engineer and AI Architect"),
    ("Email", "flourisholaiya@gmail.com"),
    ("Portfolio", "https://flourish-webpage.vercel.app/"),
    ("Repository", "github.com/lourish789/crispy-funicular"),
]:
    cells = t.add_row().cells
    r1 = cells[0].paragraphs[0].add_run(k)
    r1.bold = True
    r1.font.size = Pt(9.5)
    r1.font.color.rgb = BLACK
    r2 = cells[1].paragraphs[0].add_run(v)
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = BLACK
    cells[0].width = Inches(1.2)
    cells[1].width = Inches(5.4)

doc.save(OUT)
print("Saved:", OUT)

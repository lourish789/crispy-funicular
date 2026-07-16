"""Build the AgriTech Suite project description document (.docx).

Strictly monochrome: black text, black table rules, no colour, no emoji.
"""
from __future__ import annotations

import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
BLACK = RGBColor(0, 0, 0)
OUT = os.path.join(HERE, "AgriTech_Suite_Project_Description.docx")

doc = Document()

# ---------------------------------------------------------------- base styles
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = BLACK
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.15

for i in range(0, 5):
    try:
        st = doc.styles[f"Heading {i}"] if i else doc.styles["Title"]
    except KeyError:
        continue
    st.font.color.rgb = BLACK
    st.font.name = "Calibri"
    if i == 1:
        st.font.size = Pt(16)
    elif i == 2:
        st.font.size = Pt(13)
    elif i == 3:
        st.font.size = Pt(11.5)


def h(text, level=1, page_break=False):
    if page_break:
        doc.add_page_break()
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = BLACK
    return p


def para(text, italic=False, size=11, align=None, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = BLACK
    if align is not None:
        p.alignment = align
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.color.rgb = BLACK
    r2 = p.add_run(text)
    r2.font.color.rgb = BLACK
    return p


def _shade(cell, hexval="D9D9D9"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexval)
    tcPr.append(shd)


def table(headers, rows, widths=None, font=9.5, shade_header=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(htext)
        r.bold = True
        r.font.size = Pt(font)
        r.font.color.rgb = BLACK
        if shade_header:
            _shade(hdr[i])
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(font)
            r.font.color.rgb = BLACK
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def figure(name, caption, width=6.3):
    doc.add_picture(os.path.join(HERE, name), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = BLACK


# ================================================================ TITLE PAGE
for _ in range(4):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("AGRITECH SUITE")
r.bold = True
r.font.size = Pt(30)
r.font.color.rgb = BLACK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("A Unified Digital Platform for African Farmers and Agricultural Traders")
r.font.size = Pt(14)
r.font.color.rgb = BLACK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("_______________________________________________")
r.font.color.rgb = BLACK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Project Description Document")
r.font.size = Pt(13)
r.italic = True
r.font.color.rgb = BLACK

for _ in range(8):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Prepared by Flourish Olaiya\nJuly 2026\nVersion 1.0")
r.font.size = Pt(11)
r.font.color.rgb = BLACK

doc.add_page_break()

# ================================================================ CONTENTS
h("Contents", 1)
toc = [
    "1. Executive Summary",
    "2. About the Project",
    "3. The Problem and the People Affected",
    "4. Who the Platform Is For",
    "5. How the System Works",
    "6. How AgriTech Suite Compares with Existing Solutions",
    "7. Advantages to the User",
    "8. Designing for Everyone: Inclusivity",
    "9. Monetisation",
    "10. What Has Been Built",
    "11. Remaining Stages of Development",
    "12. Recommended Upgrades and Future Direction",
    "13. Risks and Mitigations",
    "14. Closing Summary",
    "15. Contact",
]
for item in toc:
    p = doc.add_paragraph()
    r = p.add_run(item)
    r.font.size = Pt(11)
    r.font.color.rgb = BLACK

# ================================================================ 1
h("1. Executive Summary", 1, page_break=True)
para(
    "AgriTech Suite is a single digital platform that brings together the tools an African "
    "farmer or agricultural trader needs to run a productive, profitable enterprise. Today "
    "those tools are scattered: one app diagnoses crop disease, another sells produce, a "
    "third offers advice, and most offer only one of the three. The farmer is left to stitch "
    "them together, and in practice, most never do."
)
para(
    "The platform combines six capabilities in one place: artificial-intelligence disease "
    "diagnosis for both plants and animals from a photograph or short video; a personalised "
    "advisory engine that tells a farmer what to expand, optimise, improve and abandon; a "
    "real-time agricultural news service tuned to the user's location; a marketplace that "
    "connects producers directly to buyers; a community forum for peer knowledge; and a "
    "conversational AI assistant that understands the user's specific farm."
)
para(
    "What makes the system distinctive is not any one feature but the combination, and the "
    "decision to run the disease-detection model on the device itself using open-source "
    "technology. That choice means diagnosis does not depend on a paid third-party service, "
    "does not require a strong internet connection, and costs nothing per use. It is the "
    "difference between a tool a subsistence farmer can actually afford and one they cannot."
)
para(
    "The working software exists. The backend, the frontend, the AI pipeline, the role system "
    "and the database are built and tested. The platform has not yet been deployed to public "
    "hosting; that is the immediate next step, followed by a field pilot in Nigeria and then "
    "the inclusion layer (USSD, SMS and local languages) that will extend the platform beyond "
    "smartphone owners to the majority of African farmers."
)

# ================================================================ 2
h("2. About the Project", 1, page_break=True)
h("2.1 What AgriTech Suite Is", 2)
para(
    "AgriTech Suite is best described as a one-stop agricultural operating system. A farmer "
    "signs in and finds, in one place, the answer to the four questions that determine whether "
    "a season succeeds: What is wrong with my crop or animal? What should I do differently? "
    "Who will buy what I grow, and at what price? And who can I ask when I am unsure?"
)
para(
    "The platform does not assume the user is an expert, wealthy, or permanently online. It is "
    "written for a farmer holding a phone in a field, possibly with a weak signal, who needs a "
    "clear answer in plain language rather than a dashboard of statistics."
)

h("2.2 Project Goals", 2)
table(
    ["#", "Goal", "What success looks like"],
    [
        ["1", "Put expert diagnosis in every farmer's hand",
         "A farmer identifies a crop or livestock disease in under a minute from a photo, without needing an extension officer to visit."],
        ["2", "Replace guesswork with personalised guidance",
         "Each farmer receives specific, data-driven recommendations based on their own farm profile and history, not generic advice."],
        ["3", "Shorten the distance between farmer and buyer",
         "Producers sell closer to the real market price by reaching buyers directly, rather than accepting whatever a middleman offers at the farm gate."],
        ["4", "Reduce avoidable losses",
         "Earlier disease detection and better market timing cut the crop and income a farmer loses each season."],
        ["5", "Make agricultural knowledge circulate",
         "Farmers and traders exchange practical, local knowledge in a moderated community rather than relying solely on institutions."],
        ["6", "Include the farmer who cannot afford to be included",
         "The core of the platform remains free, and works on low-end devices and weak networks."],
    ],
    widths=[0.4, 2.0, 4.0],
)

# ================================================================ 3
h("3. The Problem and the People Affected", 1, page_break=True)
para(
    "African agriculture does not suffer from a shortage of effort. It suffers from a shortage "
    "of information reaching the person holding the hoe, and a shortage of fair routes to "
    "market once the harvest is in. Four problems recur, and they compound one another."
)

h("3.1 The advice never arrives", 2)
para(
    "Agricultural extension officers are the traditional channel through which government "
    "knowledge reaches farmers. The Food and Agriculture Organization recommends roughly one "
    "extension agent per 1,000 farmers. In Nigeria the real ratio is between one agent for "
    "every 2,500 farm families and one for every 10,000, and in some states a single agent is "
    "reported to serve more than 24,000 farmers. In practice this means most farmers will never "
    "meet an extension officer in a given season. When a disease appears, they guess, they ask "
    "a neighbour who is also guessing, or they buy whatever chemical a shopkeeper recommends."
)

h("3.2 The harvest is lost after it is grown", 2)
para(
    "Nigeria lost an estimated 30 to 40 million metric tonnes of food in 2025, with the "
    "economic cost put at between 3.5 and 5 trillion naira. For perishables such as tomatoes, "
    "onions and bananas, losses reach as high as 50 per cent. A farmer can do everything right "
    "in the field and still lose half the value of the crop between harvest and sale, because "
    "there is no storage, no buyer ready, and no information about where demand actually is."
)

h("3.3 The middle takes the margin", 2)
para(
    "Because farmers rarely know the prevailing market price and cannot easily reach distant "
    "buyers, they sell at the farm gate to intermediaries. Levies imposed along the route by "
    "middlemen, transport operators and informal collectors inflate the final price paid by "
    "consumers without that value returning to the farmer. The result is the worst of both "
    "worlds: the producer earns too little and the consumer pays too much."
)

h("3.4 Technology exists but does not fit", 2)
para(
    "Nigeria has a substantial agritech sector, with roughly 145 active companies. Yet adoption "
    "among ordinary smallholders remains thin. The tools are often single-purpose, assume a "
    "smartphone and steady data, are priced for commercial farms, or are built around investors "
    "and lenders rather than the farmer's daily decisions."
)

h("3.5 Who is affected", 2)
table(
    ["Group", "How the problem shows up for them", "Scale of impact"],
    [
        ["Smallholder farmers (the majority)",
         "Crop and livestock disease diagnosed too late or wrongly; no reliable advice; forced to sell cheap at the farm gate.",
         "Highest. This group carries almost all of the loss and has the fewest alternatives."],
        ["Women farmers",
         "Lower access to extension services, credit, formal markets and, often, to smartphones and literacy in English.",
         "High and frequently overlooked."],
        ["Young and new farmers",
         "No inherited experience network; most exposed to avoidable, repeated mistakes.",
         "High; also the group most able to adopt digital tools."],
        ["Livestock keepers",
         "Veterinary care is scarcer than crop extension; disease can wipe out a herd, which is the household's stored savings.",
         "Severe when it occurs; often catastrophic."],
        ["Traders and buyers",
         "Cannot find consistent, verified supply; rely on the same opaque intermediary chain.",
         "Moderate to high; margins are eaten by inefficiency."],
        ["Processors and off-takers",
         "Unreliable volume and quality make planning and contracts difficult.",
         "Moderate; constrains industrial growth."],
        ["Consumers and the wider economy",
         "Higher food prices, unstable supply, and pressure on food security.",
         "National."],
    ],
    widths=[1.5, 3.2, 1.7],
)

# ================================================================ 4
h("4. Who the Platform Is For", 1, page_break=True)
para(
    "The primary users are Nigerian farmers and agricultural traders, with the platform "
    "designed from the outset to extend across West Africa and then the wider continent. "
    "Nigeria is the starting point because the problem is large, the mobile network is broad, "
    "and the agricultural base is enormous and largely informal. Nothing in the system is "
    "Nigeria-specific: the crops, languages, currencies and news sources are configurable."
)

h("4.1 The two roles", 2)
para(
    "Every account on the platform is one of two roles, chosen at sign-up and changeable later. "
    "This single decision shapes the whole experience."
)
table(
    ["Role", "Who they are", "What the platform gives them"],
    [
        ["Farmer (Producer)",
         "Smallholders, mixed crop-and-livestock households, commercial growers, cooperatives.",
         "Disease diagnosis, personalised advisory, the ability to publish produce for sale, and a feed showing which buyers are currently seeking what they grow."],
        ["Buyer (Consumer)",
         "Market traders, aggregators, processors, restaurants, exporters, households buying in bulk.",
         "The ability to publish buy requests, and a feed of fresh produce offered directly by farmers, with location and price."],
    ],
    widths=[1.2, 2.3, 2.9],
)

h("4.2 Representative users", 2)
table(
    ["Persona", "Situation", "What changes for them"],
    [
        ["Amina, 34 — smallholder, Kano",
         "Two hectares of tomato and maize. Basic Android phone, data is expensive. Has never met an extension officer.",
         "Photographs a diseased leaf and gets a named diagnosis and treatment before the infection spreads across the plot."],
        ["Emeka, 27 — new farmer, Enugu",
         "Inherited land, first cassava season, learning by trial and error.",
         "The advisory engine and AI assistant act as the mentor he does not have; the community answers what the machine cannot."],
        ["Bola, 45 — market trader, Lagos",
         "Buys produce weekly, depends on intermediaries, quality is inconsistent.",
         "Posts a buy request and is matched directly with farmers holding stock, with location and price visible."],
        ["Hauwa, 52 — livestock keeper, Kaduna",
         "Small cattle herd; the nearest vet is hours away and costly.",
         "Photographs a skin lesion and receives an early indication plus containment guidance while arranging veterinary help."],
        ["Chidi, 38 — processor, Aba",
         "Needs predictable volumes of raw material.",
         "Uses the platform to source consistently from a known pool of producers rather than the spot market."],
    ],
    widths=[1.5, 2.4, 2.5],
)

# ================================================================ 5
h("5. How the System Works", 1, page_break=True)
para(
    "The platform is organised in layers. The user interacts with a web application; that "
    "application talks to a single application programming interface; the interface calls the "
    "intelligence services that do the thinking; and everything is recorded in the data layer "
    "so that the system learns the user's context over time."
)
figure("fig1_architecture.png", "Figure 1: System architecture, from the user down to the data layer.")

h("5.1 The six modules", 2)
table(
    ["Module", "What it does for the user", "How it works underneath"],
    [
        ["Disease diagnosis",
         "Identifies disease in a crop or animal from a photo or short video and returns the disease name, its cause, an immediate solution and prevention steps.",
         "An open-source computer-vision model runs on the server/device and classifies the image; a language model turns that label into practical advice. Videos are sampled into frames first."],
        ["Advisory engine",
         "Tells the farmer what to expand, what to optimise, what to improve, and what to give up on.",
         "Combines the user's profile and their real history on the platform with rule-based analysis and a language model."],
        ["Agri-news",
         "Global and local agricultural news relevant to the user's location.",
         "Aggregates news sources and caches results so it is fast and cheap to serve."],
        ["Marketplace",
         "Connects producers and buyers directly.",
         "Listings are typed automatically by the user's role, and each side is shown the opposite side's activity."],
        ["Community",
         "A forum to ask, answer and share by topic.",
         "Posts, comments and likes, each tagged with the author's role so users know whether advice comes from a farmer or a buyer."],
        ["AI assistant",
         "A conversational agricultural expert that already knows the user's farm.",
         "A language model with the user's profile, history and a searchable knowledge base injected into every conversation, streamed word-by-word."],
    ],
    widths=[1.1, 2.6, 2.7],
)

h("5.2 How a diagnosis actually happens", 2)
para(
    "This is the flow a farmer experiences most often, and it was designed so that a failure "
    "anywhere in the chain still produces a useful answer rather than an error message."
)
figure("fig2_diagnosis_flow.png", "Figure 2: The diagnosis pipeline, including its fallback path.")
para(
    "The important detail is the fallback chain. If the on-device model is unavailable, the "
    "system tries a cloud vision service, then reasoning over the image's measurable "
    "characteristics, and finally a rule-based assessment. The farmer always receives a "
    "structured answer. A tool that fails silently in a field with no signal is not a tool.",
    italic=False,
)

h("5.3 How the marketplace connects the two sides", 2)
para(
    "The marketplace is deliberately two-sided and role-aware. A farmer never has to understand "
    "listing categories; they simply post what they have, and the system publishes it as a sell "
    "offer. A buyer posts what they need, and it becomes a buy request. Each is then shown the "
    "other side."
)
figure("fig3_marketplace.png", "Figure 3: Role-based matching between producers and buyers.")

# ================================================================ 6
h("6. How AgriTech Suite Compares with Existing Solutions", 1, page_break=True)
para(
    "The comparison below is drawn from published information about established platforms "
    "operating in Nigeria and across Africa. The purpose is not to diminish them: several are "
    "excellent and have reached far more farmers than this project currently has. The purpose "
    "is to be precise about the gap AgriTech Suite is built to fill."
)

h("6.1 The existing landscape", 2)
table(
    ["Platform", "Core focus", "Strengths", "Limitations relative to this project"],
    [
        ["PlantVillage Nuru",
         "Offline AI crop-disease diagnosis (Penn State / USAID).",
         "Excellent, field-proven accuracy; works offline; local languages; strong research base and extension-worker network.",
         "Diagnosis only. No marketplace, no trading, no personalised business advisory, no livestock coverage."],
        ["Plantix",
         "Crop-disease diagnosis and community.",
         "Very large image library; wide crop coverage; community features.",
         "Advice is generic rather than tied to the individual farm's history; no integrated route to market."],
        ["Farmcrowdy",
         "Crowdfunding of farm cycles.",
         "Pioneered digital agriculture financing in Nigeria; connects sponsors to farms.",
         "Built primarily around the investor relationship, not the farmer's daily agronomic decisions."],
        ["Thrive Agric",
         "Finance, inputs and technical support.",
         "Real access to credit and discounted inputs; strong farmer network.",
         "Credit-led. Does not provide AI diagnosis, and the farmer's market access is mediated by the programme."],
        ["AFEX",
         "Commodities exchange and warehousing.",
         "Transparent pricing; genuine market infrastructure; storage.",
         "Serves the formal commodity trade; less suited to the individual smallholder with one crate of tomatoes."],
        ["Crop2Cash",
         "Financial inclusion via USSD.",
         "Reaches feature-phone farmers; over 500,000 farmers; payments, credit, insurance.",
         "Financial services rather than agronomy; no disease diagnosis or advisory intelligence."],
        ["Releaf",
         "Crop processing technology.",
         "Adds real value in the processing chain.",
         "Focused on processing, not on the farmer's information problem."],
        ["AgriTech Suite",
         "All six: diagnosis, advisory, news, marketplace, community, AI assistant.",
         "Only platform combining on-device open-source diagnosis for both plants and animals with a role-based direct marketplace and per-farm personalised advisory.",
         "Earliest stage: not yet deployed, not yet field-tested at scale, no user base or financing rails yet."],
    ],
    widths=[0.95, 1.25, 2.1, 2.1],
    font=8.5,
)

h("6.2 Capability matrix", 2)
para("A direct comparison of what each platform does and does not do.")
table(
    ["Capability", "Nuru", "Plantix", "Farmcrowdy", "Thrive Agric", "AFEX", "Crop2Cash", "AgriTech Suite"],
    [
        ["Plant disease diagnosis", "Yes", "Yes", "No", "Partial", "No", "No", "Yes"],
        ["Animal disease diagnosis", "No", "No", "No", "No", "No", "No", "Yes"],
        ["Works without a paid AI service", "Yes", "No", "n/a", "n/a", "n/a", "n/a", "Yes"],
        ["Personalised advisory (per farm)", "Partial", "No", "No", "Partial", "No", "No", "Yes"],
        ["Direct farmer-to-buyer marketplace", "No", "No", "Partial", "Partial", "Yes", "No", "Yes"],
        ["Buyer role with matched feed", "No", "No", "No", "No", "Partial", "No", "Yes"],
        ["Community forum", "No", "Yes", "No", "No", "No", "No", "Yes"],
        ["Conversational AI assistant", "No", "No", "No", "No", "No", "No", "Yes"],
        ["Localised agri-news", "No", "No", "No", "No", "Partial", "No", "Yes"],
        ["Feature phone (USSD/SMS)", "No", "No", "No", "Partial", "No", "Yes", "Planned"],
        ["Full offline operation", "Yes", "Partial", "No", "No", "No", "Yes", "Planned"],
        ["Deployed and in the field today", "Yes", "Yes", "Yes", "Yes", "Yes", "Yes", "Not yet"],
    ],
    widths=[1.85, 0.62, 0.68, 0.85, 0.85, 0.6, 0.78, 0.95],
    font=8.5,
)

h("6.3 What the comparison shows", 2)
para(
    "Two conclusions follow honestly from the table. The first is that the market is not short "
    "of point solutions: diagnosis is solved well by Nuru, financial inclusion by Crop2Cash, "
    "commodity trading by AFEX. What no one has assembled is the whole picture for one farmer, "
    "in one place, where the diagnosis feeds the advisory, and the advisory feeds the decision "
    "about what to plant and where to sell it."
)
para(
    "The second conclusion is that this project is behind on the things that only time and users "
    "produce. Nuru is field-proven with tens of thousands of farmers and local-language support. "
    "Crop2Cash reaches feature phones today. AgriTech Suite is technically complete but "
    "unproven in the field. That is the honest position, and it defines the roadmap in Section 11: "
    "deploy, pilot, then close the inclusion gap."
)

# ================================================================ 7
h("7. Advantages to the User", 1, page_break=True)
para(
    "A platform is only worth a farmer's attention if it changes something measurable in their "
    "season. The table below states the advantage and, importantly, the mechanism by which it "
    "is delivered, so the claim can be tested rather than taken on faith."
)
table(
    ["Advantage", "Why it matters to the user", "How the platform delivers it"],
    [
        ["Expert diagnosis without an expert",
         "Where one extension officer serves thousands, waiting for a visit means losing the crop.",
         "On-device AI returns a named disease and a treatment plan in seconds, at any hour."],
        ["Costs nothing per diagnosis",
         "A per-use fee prices out the very farmers who most need the tool.",
         "The vision model is open-source and runs locally, so the marginal cost of a diagnosis is effectively zero."],
        ["Works when the network does not",
         "Rural connectivity is intermittent; a tool that needs strong data fails exactly where it is needed.",
         "Local model plus an offline knowledge base and a graceful fallback chain."],
        ["Advice about this farm, not farms in general",
         "Generic advice is easy to ignore because it is not about your soil, your crop or your history.",
         "The advisory engine reads the user's own profile and platform history before recommending anything."],
        ["Covers animals as well as crops",
         "Livestock is often a household's savings; veterinary access is even thinner than crop extension.",
         "A separate open-source model handles livestock conditions, with the same structured output."],
        ["A better price for the harvest",
         "Farm-gate selling to intermediaries is where much of the farmer's margin disappears.",
         "Direct listings and a buyer feed put the producer in contact with demand, with prices visible."],
        ["Demand visibility before planting",
         "Growing what nobody is buying is a season wasted.",
         "The farmer's feed shows live buy requests, which informs what to plant and when to harvest."],
        ["Knowledge that compounds",
         "A single answer helps once; a record builds judgement.",
         "Every diagnosis and interaction is stored and feeds future advisory and AI conversations."],
        ["Plain language, on demand",
         "Technical agronomy is a barrier for many users.",
         "The AI assistant answers conversationally and is designed to be patient and non-judgemental."],
        ["Peer knowledge with context",
         "Farmers trust other farmers.",
         "The community shows each author's role, so users know whether they are hearing from a producer or a buyer."],
    ],
    widths=[1.5, 2.4, 2.5],
)

# ================================================================ 8
h("8. Designing for Everyone: Inclusivity", 1, page_break=True)
para(
    "The users of this platform will not be alike. Some will farm hundreds of hectares and run "
    "the business from a laptop. Many more will have a shared phone, an intermittent signal, a "
    "prepaid balance measured in hundreds of naira, and limited reading confidence in English. "
    "A platform that only serves the first group is not an agricultural platform; it is a "
    "business tool for people who already have options."
)
para(
    "The design principle is therefore that capability should degrade gracefully with the "
    "user's means, never disappear. The same underlying intelligence should be reachable "
    "through four progressively simpler doors."
)
figure("fig4_inclusivity.png", "Figure 4: Four routes into the same platform, by device and means.")

h("8.1 Commitments that make inclusion real", 2)
table(
    ["Commitment", "What it means in practice"],
    [
        ["The core is free, permanently",
         "Diagnosis, advisory, news, community and basic listings are free for every user. Revenue comes from transactions and optional services, never from access to knowledge."],
        ["No smartphone requirement",
         "A planned USSD short-code and SMS service give feature-phone users price alerts, buy/sell matching and advisory summaries without any data bundle."],
        ["Low data by default",
         "On-device processing, aggressive caching and small payloads. A farmer should not have to choose between a diagnosis and airtime."],
        ["Local languages and voice",
         "Planned support for Hausa, Yoruba, Igbo and Nigerian Pidgin, then Swahili and French, with voice input and spoken answers for users with low literacy."],
        ["Designed for the shared phone",
         "Simple sign-in, low friction, and no assumption that the device belongs to one person."],
        ["Gender-aware access",
         "Because women farmers face lower access to extension, credit and devices, the agent and cooperative model is intended to reach them where the app alone would not."],
        ["No penalty for being poor",
         "The free tier is not a crippled demonstration. It is the whole agronomic product. Paid tiers add commercial reach, not knowledge."],
    ],
    widths=[1.7, 4.7],
)

# ================================================================ 9
h("9. Monetisation", 1, page_break=True)
para(
    "The commercial design follows one rule: the platform should earn when the user earns, and "
    "not before. A model that charges a subsistence farmer for advice will be abandoned, and "
    "deserves to be. A model that takes a small share of a sale the farmer would not otherwise "
    "have made is one the farmer will gladly pay."
)

h("9.1 How the user makes and saves money", 2)
para(
    "This is the part that matters most to the user, and it is worth stating before the "
    "platform's own revenue."
)
table(
    ["Route", "How the user gains"],
    [
        ["Selling closer to the real price",
         "Reaching buyers directly rather than accepting a farm-gate offer recovers margin that currently goes to intermediaries."],
        ["Losing less of the harvest",
         "Earlier disease detection and better timing directly reduce the share of the crop lost — the single largest recoverable loss in the chain."],
        ["Spending less on the wrong inputs",
         "A correct diagnosis prevents buying and spraying the wrong chemical, which is both wasted money and wasted crop."],
        ["Planting what is actually in demand",
         "Visibility of live buy requests reduces the risk of producing a crop with no buyer."],
        ["Selling knowledge and reputation",
         "Experienced farmers can build a verified reputation in the community and marketplace, attracting better buyers and, in future, paid advisory or agent roles."],
        ["Becoming an agent",
         "The planned agent model lets a trusted local user serve neighbours who lack devices, earning a commission — turning inclusion into a livelihood."],
        ["Building a financial record",
         "A history of verified sales on the platform becomes evidence of creditworthiness for farmers who currently have none."],
    ],
    widths=[1.7, 4.7],
)

h("9.2 How the platform earns", 2)
figure("fig6_revenue.png", "Figure 5: Revenue streams, sitting on top of a permanently free core.")
table(
    ["Stream", "Model", "Who pays", "Rationale"],
    [
        ["Marketplace commission",
         "A small percentage of a completed transaction.",
         "Buyer or seller on a successful trade.",
         "The primary stream. It is aligned: no sale, no fee. Comparable African platforms charge in the region of 5-15% depending on the service."],
        ["Premium subscription",
         "Low monthly fee.",
         "Active traders and commercial farmers.",
         "Verified badge, priority placement, price alerts and advanced buyer matching. Optional, never required to use the core."],
        ["Business / API tier",
         "Annual contract.",
         "Cooperatives, processors, exporters, NGOs.",
         "Bulk sourcing, analytics and integration. These users have budgets and derive clear operational value."],
        ["Input and service referrals",
         "Referral or affiliate fee.",
         "Seed, fertiliser, veterinary and logistics providers.",
         "The diagnosis already identifies the need; connecting it to a vetted supplier serves the farmer and monetises the moment."],
        ["Financing and insurance partnerships",
         "Revenue share.",
         "Lenders and insurers.",
         "Platform history de-risks lending. The farmer gains access to credit they could not otherwise obtain."],
        ["Aggregate insight reports",
         "Licensed reports.",
         "Government, research bodies, agribusiness.",
         "Anonymised and consented only. Disease-outbreak and price-trend data has genuine public value."],
        ["Grants and development funding",
         "Non-dilutive.",
         "Development agencies.",
         "Appropriate for the inclusion layer (USSD, languages), which is a public good and hard to fund commercially."],
    ],
    widths=[1.2, 1.1, 1.4, 2.7],
    font=8.5,
)

h("9.3 Tiers", 2)
table(
    ["Tier", "Price", "Who it is for", "What is included"],
    [
        ["Free", "Zero, permanently", "Every farmer; the majority of users",
         "Diagnosis (plant and animal), advisory, news, community, AI assistant, basic listings."],
        ["Premium", "Low monthly", "Active sellers and buyers",
         "Everything in Free, plus verified badge, priority listing, price alerts, advanced matching."],
        ["Business", "Annual", "Cooperatives, processors, exporters",
         "Everything in Premium, plus bulk tools, analytics, multi-user accounts, API access."],
        ["Agent", "Earns commission", "Trusted local users serving offline neighbours",
         "Tools to transact on behalf of others; a route to income and to inclusion."],
    ],
    widths=[0.85, 1.05, 1.7, 2.8],
)
para(
    "The sequencing matters. Commission on trade should be the first stream switched on, because "
    "it requires no behaviour change and no upfront payment. Subscriptions follow once the "
    "marketplace is liquid enough that priority placement is worth paying for. Data and "
    "financing partnerships come last, once volume makes them credible.",
    italic=True,
)

# ================================================================ 10
h("10. What Has Been Built", 1, page_break=True)
para(
    "The following is an accurate statement of the current build. Everything marked complete "
    "exists in the repository, runs, and is covered by an automated test suite that passes."
)
table(
    ["Component", "Status", "Detail"],
    [
        ["Backend API", "Complete",
         "A REST interface covering authentication, diagnosis, advisory, news, marketplace, community and chat."],
        ["Database schema", "Complete",
         "Users with farm profiles and roles, diagnoses, listings, posts, comments, chat history, and a vector knowledge store. Includes safe, additive migrations."],
        ["Plant disease diagnosis", "Complete",
         "Open-source model with 38 disease classes, verified returning real classifications on test images."],
        ["Animal disease diagnosis", "Complete",
         "A separate open-source model for livestock skin conditions, sharing the same structured output."],
        ["Image and video handling", "Complete",
         "Images accepted directly; videos sampled into representative frames automatically. Verified on a real video file."],
        ["Advisory engine", "Complete",
         "Produces expand / optimise / improve / give-up guidance from the user's own profile and history."],
        ["Agri-news aggregator", "Complete",
         "Global and location-specific news with time-based caching."],
        ["Marketplace", "Complete",
         "Listings, search, categories, and role-based matching between producers and buyers."],
        ["Community", "Complete",
         "Posts, comments, likes, topic filtering, and author role badges."],
        ["AI assistant", "Complete",
         "Conversational assistant with the user's profile, history and a retrieval knowledge base injected into context, streaming replies word-by-word."],
        ["Roles (Farmer / Buyer)", "Complete",
         "Selected at sign-up, changeable in profile, and enforced across the marketplace and community."],
        ["Authentication", "Complete",
         "Email and password with secure tokens, plus Firebase sign-in integration (requires project credentials to activate)."],
        ["Web frontend", "Complete",
         "Nine screens covering sign-in, dashboard, diagnosis, advisory, news, marketplace, community, chat and profile."],
        ["Resilience", "Complete",
         "Every external AI or news call has a deterministic fallback; the platform degrades rather than fails."],
        ["Automated tests", "Complete",
         "An end-to-end suite covering every endpoint, all passing."],
        ["Public deployment", "Not started",
         "The backend is not yet on Render and the frontend is not yet on Vercel. The application currently runs locally only."],
        ["Field validation", "Not started",
         "No pilot has been run with real farmers; accuracy in real field conditions is not yet measured."],
        ["USSD / SMS access", "Not started", "Planned for the inclusion phase."],
        ["Local languages", "Not started", "Planned for the inclusion phase."],
        ["Payments and escrow", "Not started", "Planned once the marketplace has real liquidity."],
    ],
    widths=[1.5, 0.95, 3.95],
    font=9,
)

# ================================================================ 11
h("11. Remaining Stages of Development", 1, page_break=True)
figure("fig5_roadmap.png", "Figure 6: Development roadmap. Phase 1 is complete; Phase 2 is the immediate next step.")

h("11.1 Phase 2 — Deployment and pilot (immediate)", 2)
para(
    "The single most important next step is to put the platform in front of real users. The "
    "software is finished but has never met a farmer, and no amount of further engineering "
    "substitutes for that."
)
table(
    ["Step", "Work involved"],
    [
        ["Deploy the backend to Render",
         "Provision the service, move from the local database to a managed PostgreSQL instance, configure environment variables and API keys, and enable health monitoring."],
        ["Deploy the frontend to Vercel",
         "Connect the repository, set the API base URL and any Firebase keys, and configure the production domain."],
        ["Harden for production",
         "Restrict cross-origin access to the real domain, add rate limiting, move uploaded media to object storage, and add error tracking."],
        ["Activate Firebase",
         "Create the Firebase project, enable Google sign-in, and supply the service credentials so federated sign-in becomes live."],
        ["Run a controlled pilot",
         "Recruit a small cohort of farmers and buyers in one Nigerian state. Measure diagnosis accuracy against expert assessment, and record what users actually do rather than what they say."],
        ["Measure and correct",
         "Use pilot data to retrain or re-select the vision models and to rewrite advisory content that does not land."],
    ],
    widths=[1.7, 4.7],
)

h("11.2 Phase 3 — Inclusion", 2)
table(
    ["Step", "Work involved"],
    [
        ["USSD short-code and SMS", "Partner with a telecommunications aggregator so feature-phone users can diagnose by SMS description, receive price alerts, and post or answer listings without data."],
        ["Offline-first mobile application", "A progressive web app or Android build with the vision model compressed for on-device use, syncing when a connection returns."],
        ["Local languages", "Hausa, Yoruba, Igbo and Nigerian Pidgin first, then Swahili and French as the platform moves beyond Nigeria."],
        ["Voice interface", "Spoken questions and spoken answers, for users with limited literacy."],
        ["Agent network", "Recruit and equip trusted local agents to serve farmers without devices, earning commission."],
    ],
    widths=[1.7, 4.7],
)

h("11.3 Phase 4 — Commerce and scale", 2)
table(
    ["Step", "Work involved"],
    [
        ["Payments and escrow", "Integrate a Nigerian payment provider so trades settle on-platform, with escrow to build trust between strangers."],
        ["Logistics", "Partner for transport and aggregation so a matched trade can actually move."],
        ["Reputation and verification", "Ratings, verified sellers and dispute resolution."],
        ["Credit and insurance", "Use verified platform history to unlock lending and crop insurance for farmers with no formal record."],
        ["Weather and satellite data", "Add forecasts and vegetation indices to make advisory predictive rather than reactive."],
        ["Regional expansion", "Ghana, Kenya and francophone West Africa, following the language work."],
    ],
    widths=[1.7, 4.7],
)

# ================================================================ 12
h("12. Recommended Upgrades and Future Direction", 1, page_break=True)
para(
    "The following ideas emerged from reviewing what comparable platforms across Africa have "
    "learned. They are ordered by the ratio of impact to effort, and each is offered with the "
    "reasoning behind it."
)
table(
    ["#", "Upgrade", "Why it is worth doing"],
    [
        ["1", "Lead with USSD, not the app store",
         "Evidence from across Africa is consistent: USSD works on any handset, needs no data bundle and costs the user almost nothing. One Acre Fund enrolled over 600,000 farmers in Rwanda through USSD alone. This is the single highest-leverage addition to the platform."],
        ["2", "A WhatsApp assistant",
         "WhatsApp is already where Nigerian farmers and traders talk. Meeting them there removes the need to persuade anyone to install anything, and the AI assistant transfers to it almost directly."],
        ["3", "Compress the vision model for on-device use",
         "Converting the model to a mobile-optimised format would let diagnosis run entirely on the handset with no upload at all — faster, free, and fully offline."],
        ["4", "Expand the livestock model",
         "Current livestock coverage is cattle skin conditions. Poultry disease is arguably the higher-frequency need in Nigeria and would broaden the platform's reach substantially."],
        ["5", "Collect field images to build a proprietary dataset",
         "Every farmer photo, with consent, improves the model. Over time a locally-trained dataset of African crops in African conditions becomes the platform's most defensible asset — this is precisely how PlantVillage built its accuracy."],
        ["6", "Price intelligence and forecasting",
         "Telling a farmer that the price of tomatoes will likely rise in ten days changes the harvest decision. This turns the platform from informative to decisive."],
        ["7", "Escrow-backed payments",
         "The reason farmers and buyers still use intermediaries is trust. Escrow manufactures trust between strangers and is what makes disintermediation actually happen."],
        ["8", "Cooperative and group accounts",
         "Most smallholders already organise into cooperatives. Serving the group rather than only the individual is the fastest realistic route to scale."],
        ["9", "Outbreak mapping",
         "Aggregating anonymised diagnoses reveals disease outbreaks geographically. This is valuable to farmers as an early warning and to government as public-health infrastructure."],
        ["10", "Traceability for export",
         "A verified record from farm to buyer supports export compliance and premium pricing — a route to significantly higher farmer income."],
        ["11", "Human expert escalation",
         "When AI confidence is low, route the case to a real agronomist or vet. This protects users from confident wrong answers and creates a paid premium service."],
        ["12", "Formal accuracy benchmarking",
         "Publish measured accuracy against expert diagnosis. Credibility with farmers, partners and funders depends on evidence, not claims."],
    ],
    widths=[0.35, 1.55, 4.5],
    font=9,
)

# ================================================================ 13
h("13. Risks and Mitigations", 1, page_break=True)
table(
    ["Risk", "Assessment", "Mitigation"],
    [
        ["Model accuracy fails in real field conditions",
         "The models are trained largely on curated datasets. Real photographs are blurry, badly lit and cluttered.",
         "Measure against expert diagnosis during the pilot; show confidence scores honestly; escalate low-confidence cases to a human; retrain on collected field images."],
        ["Farmers do not adopt the platform",
         "The most common failure of agritech. A good tool nobody uses is worth nothing.",
         "Lead with the free diagnosis, which delivers value on first use; reach users through cooperatives and agents; add USSD so a smartphone is not required."],
        ["Marketplace has no liquidity",
         "A two-sided market is worthless until both sides show up.",
         "Launch in one state with one crop to concentrate density rather than spreading thin; recruit buyers before farmers."],
        ["Trust between strangers",
         "Farmers and buyers use intermediaries partly because intermediaries are known to them.",
         "Escrow, verification, ratings and dispute resolution; start within existing cooperative networks where trust already exists."],
        ["Dependence on third-party AI services",
         "Providers change and retire models; one such model was already retired during this project.",
         "The primary detector is open-source and self-hosted; every provider model is configurable without code changes."],
        ["Connectivity and cost",
         "Rural data is unreliable and expensive.",
         "On-device processing, offline fallbacks, caching, and the planned USSD and SMS layer."],
        ["Wrong advice causes harm",
         "A confidently wrong diagnosis could cost a farmer a season.",
         "Show confidence; phrase advice as guidance; recommend professional confirmation for severe cases; never overstate certainty."],
        ["Sustainability before revenue",
         "The free core must be funded until commission volume matures.",
         "Deliberately low running costs through open-source models; pursue development grants for the inclusion layer, which is a public good."],
    ],
    widths=[1.3, 2.4, 2.7],
    font=8.5,
)

# ================================================================ 14
h("14. Closing Summary", 1, page_break=True)
para(
    "AgriTech Suite starts from an unglamorous observation: the African farmer's problem is "
    "usually not that the knowledge does not exist, but that it never reaches them in time, and "
    "that when the harvest finally comes, someone else captures most of its value. Extension "
    "officers are outnumbered thousands to one. Trillions of naira of food is lost every year "
    "after it has already been successfully grown. And the tools built to help are often "
    "single-purpose, priced for people who are not poor, or dependent on connectivity that "
    "rural users do not have."
)
para(
    "This project's response is to put the whole chain in one place — see the problem, "
    "understand it, act on it, sell the result, and ask someone when unsure — and to build the "
    "expensive part, the disease detection, on open-source technology that runs locally so it "
    "can be given away free forever. The platform earns from transactions the farmer would not "
    "otherwise have made, which keeps the incentives pointed in the right direction."
)
para(
    "The software is built and tested. It is not yet deployed, and it has not yet been "
    "validated with real farmers in real fields — those are the next two steps, in that order. "
    "The technical foundation is complete enough that the remaining work is about reach and "
    "evidence rather than construction: get it online, prove the accuracy honestly, and then "
    "extend it to the feature-phone majority who need it most and who no smartphone-only "
    "platform will ever serve."
)

# ================================================================ 15
h("15. Contact", 1)
doc.add_paragraph()
t = doc.add_table(rows=0, cols=2)
t.style = "Table Grid"
contact = [
    ("Author", "Flourish Olaiya"),
    ("Role", "Full-Stack Software Engineer and AI Architect"),
    ("Project", "AgriTech Suite"),
    ("Email", "flourisholaiya@gmail.com"),
    ("Portfolio", "https://flourish-webpage.vercel.app/"),
    ("Repository", "github.com/lourish789/crispy-funicular"),
    ("Document version", "1.0 — July 2026"),
]
for k, v in contact:
    cells = t.add_row().cells
    p1 = cells[0].paragraphs[0]
    r1 = p1.add_run(k)
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = BLACK
    p2 = cells[1].paragraphs[0]
    r2 = p2.add_run(v)
    r2.font.size = Pt(10)
    r2.font.color.rgb = BLACK
    cells[0].width = Inches(1.6)
    cells[1].width = Inches(4.8)

doc.add_paragraph()
para(
    "This document describes the AgriTech Suite project as built to date. Figures for "
    "post-harvest losses, extension-officer ratios and comparable platforms are drawn from "
    "published Nigerian and pan-African agricultural sources current at the time of writing.",
    italic=True, size=8.5,
)

doc.save(OUT)
print("Saved:", OUT)
